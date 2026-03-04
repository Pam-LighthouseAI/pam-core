#!/usr/bin/env python3
"""
Grant Template Library: Markdown to Word Document Converter

Converts markdown grant templates to professionally formatted Word documents (.docx)
with proper styles, headings, tables, and formatting.
"""

import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def create_styles(doc):
    """Create custom styles for the document."""
    styles = doc.styles
    
    # Title Style
    if 'GrantTitle' not in [s.name for s in styles]:
        title_style = styles.add_style('GrantTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = 'Calibri'
        title_style.font.size = Pt(24)
        title_style.font.bold = True
        title_style.font.color.rgb = RGBColor(0, 51, 102)
        title_style.paragraph_format.space_after = Pt(12)
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Subtitle Style
    if 'GrantSubtitle' not in [s.name for s in styles]:
        subtitle_style = styles.add_style('GrantSubtitle', WD_STYLE_TYPE.PARAGRAPH)
        subtitle_style.font.name = 'Calibri'
        subtitle_style.font.size = Pt(14)
        subtitle_style.font.italic = True
        subtitle_style.font.color.rgb = RGBColor(102, 102, 102)
        subtitle_style.paragraph_format.space_after = Pt(12)
    
    # Body Text Style
    if 'GrantBody' not in [s.name for s in styles]:
        body_style = styles.add_style('GrantBody', WD_STYLE_TYPE.PARAGRAPH)
        body_style.font.name = 'Calibri'
        body_style.font.size = Pt(11)
        body_style.paragraph_format.space_after = Pt(8)
        body_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    
    # Quote Style
    if 'GrantQuote' not in [s.name for s in styles]:
        quote_style = styles.add_style('GrantQuote', WD_STYLE_TYPE.PARAGRAPH)
        quote_style.font.name = 'Calibri'
        quote_style.font.size = Pt(11)
        quote_style.font.italic = True
        quote_style.font.color.rgb = RGBColor(68, 68, 68)
        quote_style.paragraph_format.left_indent = Inches(0.5)
        quote_style.paragraph_format.space_after = Pt(8)
    
    # Code Style
    if 'GrantCode' not in [s.name for s in styles]:
        code_style = styles.add_style('GrantCode', WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = 'Consolas'
        code_style.font.size = Pt(9)
        code_style.font.color.rgb = RGBColor(51, 51, 51)
        code_style.paragraph_format.left_indent = Inches(0.25)
        code_style.paragraph_format.space_after = Pt(4)
    
    # Configure Heading styles
    for i in range(1, 4):
        heading_style = styles[f'Heading {i}']
        heading_style.font.name = 'Calibri'
        heading_style.font.color.rgb = RGBColor(0, 51, 102)
        if i == 1:
            heading_style.font.size = Pt(18)
            heading_style.font.bold = True
            heading_style.paragraph_format.space_before = Pt(18)
            heading_style.paragraph_format.space_after = Pt(6)
        elif i == 2:
            heading_style.font.size = Pt(14)
            heading_style.font.bold = True
            heading_style.paragraph_format.space_before = Pt(12)
            heading_style.paragraph_format.space_after = Pt(6)
        else:
            heading_style.font.size = Pt(12)
            heading_style.font.bold = True
            heading_style.paragraph_format.space_before = Pt(10)
            heading_style.paragraph_format.space_after = Pt(4)


def parse_table(lines, start_idx):
    """Parse a markdown table and return table data and end index."""
    table_data = []
    i = start_idx
    
    while i < len(lines):
        line = lines[i].strip()
        if not line or not line.startswith('|'):
            break
        
        # Skip separator lines (|---|---|)
        if re.match(r'^\|[\s\-:|]+\|$', line):
            i += 1
            continue
        
        # Parse table row
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        if cells:
            table_data.append(cells)
        i += 1
    
    return table_data, i


def add_table(doc, table_data):
    """Add a formatted table to the document."""
    if not table_data:
        return
    
    # Create table
    num_rows = len(table_data)
    num_cols = len(table_data[0])
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Populate and format table
    for row_idx, row_data in enumerate(table_data):
        row = table.rows[row_idx]
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < len(row.cells):
                cell = row.cells[col_idx]
                cell.text = cell_text
                
                # Format header row
                if row_idx == 0:
                    set_cell_shading(cell, '003366')  # Dark blue
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(255, 255, 255)
                            run.font.size = Pt(10)
                else:
                    # Alternate row colors
                    if row_idx % 2 == 0:
                        set_cell_shading(cell, 'F2F2F2')  # Light gray
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)
                
                # Set cell padding
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(2)
                    paragraph.paragraph_format.space_after = Pt(2)
    
    # Add space after table
    doc.add_paragraph()


def parse_markdown(md_content):
    """Parse markdown content and return structured data."""
    lines = md_content.split('\n')
    elements = []
    i = 0
    
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # Empty line
        if not stripped:
            elements.append(('empty', ''))
            i += 1
            continue
        
        # Horizontal rule
        if stripped in ['---', '***', '___']:
            elements.append(('hr', ''))
            i += 1
            continue
        
        # Heading 1
        if stripped.startswith('# '):
            elements.append(('h1', stripped[2:]))
            i += 1
            continue
        
        # Heading 2
        if stripped.startswith('## '):
            elements.append(('h2', stripped[3:]))
            i += 1
            continue
        
        # Heading 3
        if stripped.startswith('### '):
            elements.append(('h3', stripped[4:]))
            i += 1
            continue
        
        # Heading 4
        if stripped.startswith('#### '):
            elements.append(('h4', stripped[5:]))
            i += 1
            continue
        
        # Table
        if stripped.startswith('|'):
            table_data, i = parse_table(lines, i)
            if table_data:
                elements.append(('table', table_data))
            continue
        
        # Code block
        if stripped.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # Skip closing ```
            elements.append(('code', '\n'.join(code_lines)))
            continue
        
        # Blockquote
        if stripped.startswith('>'):
            quote_text = stripped[1:].strip()
            elements.append(('quote', quote_text))
            i += 1
            continue
        
        # Checkbox list
        checkbox_match = re.match(r'^- \[([ x])\]\s*(.*)', stripped)
        if checkbox_match:
            checked = checkbox_match.group(1) == 'x'
            text = checkbox_match.group(2)
            elements.append(('checkbox', (checked, text)))
            i += 1
            continue
        
        # Unordered list
        if stripped.startswith('- ') or stripped.startswith('* '):
            elements.append(('bullet', stripped[2:]))
            i += 1
            continue
        
        # Ordered list
        numbered_match = re.match(r'^(\d+)\.\s*(.*)', stripped)
        if numbered_match:
            number = int(numbered_match.group(1))
            text = numbered_match.group(2)
            elements.append(('numbered', (number, text)))
            i += 1
            continue
        
        # Bold text line (like **Version:** 1.0)
        if stripped.startswith('**') and '**' in stripped[2:]:
            elements.append(('bold_line', stripped))
            i += 1
            continue
        
        # Regular paragraph
        elements.append(('paragraph', stripped))
        i += 1
    
    return elements


def format_inline_text(paragraph, text):
    """Apply inline formatting (bold, italic, code) to text."""
    # Handle bold and italic
    parts = []
    current = ''
    i = 0
    
    while i < len(text):
        # Bold italic
        if text[i:i+3] == '***':
            if current:
                parts.append(('normal', current))
                current = ''
            j = i + 3
            while j < len(text) and text[j:j+3] != '***':
                j += 1
            parts.append(('bold_italic', text[i+3:j]))
            i = j + 3
            continue
        
        # Bold
        if text[i:i+2] == '**':
            if current:
                parts.append(('normal', current))
                current = ''
            j = i + 2
            while j < len(text) and text[j:j+2] != '**':
                j += 1
            parts.append(('bold', text[i+2:j]))
            i = j + 2
            continue
        
        # Italic
        if text[i] == '*' and (i == 0 or text[i-1] != '*'):
            if current:
                parts.append(('normal', current))
                current = ''
            j = i + 1
            while j < len(text) and text[j] != '*':
                j += 1
            parts.append(('italic', text[i+1:j]))
            i = j + 1
            continue
        
        # Inline code
        if text[i] == '`':
            if current:
                parts.append(('normal', current))
                current = ''
            j = i + 1
            while j < len(text) and text[j] != '`':
                j += 1
            parts.append(('code', text[i+1:j]))
            i = j + 1
            continue
        
        current += text[i]
        i += 1
    
    if current:
        parts.append(('normal', current))
    
    # Add runs to paragraph
    for style, content in parts:
        run = paragraph.add_run(content)
        if style == 'bold':
            run.bold = True
        elif style == 'italic':
            run.italic = True
        elif style == 'bold_italic':
            run.bold = True
            run.italic = True
        elif style == 'code':
            run.font.name = 'Consolas'
            run.font.size = Pt(10)


def convert_md_to_docx(md_path, docx_path):
    """Convert a markdown file to a Word document."""
    # Read markdown content
    with open(md_path, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Create document
    doc = Document()
    
    # Set up styles
    create_styles(doc)
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Parse markdown
    elements = parse_markdown(md_content)
    
    # Track list state
    in_bullet_list = False
    in_numbered_list = False
    bullet_list = []
    numbered_list = []
    
    for elem_type, content in elements:
        # Handle lists
        if elem_type == 'bullet':
            if not in_bullet_list:
                in_bullet_list = True
                bullet_list = []
            bullet_list.append(content)
            continue
        elif in_bullet_list and elem_type not in ('bullet', 'empty'):
            # Output bullet list
            for item in bullet_list:
                p = doc.add_paragraph(style='List Bullet')
                format_inline_text(p, item)
            bullet_list = []
            in_bullet_list = False
        
        if elem_type == 'numbered':
            if not in_numbered_list:
                in_numbered_list = True
                numbered_list = []
            numbered_list.append(content[1])  # Just the text, not the number
            continue
        elif in_numbered_list and elem_type not in ('numbered', 'empty'):
            # Output numbered list
            for item in numbered_list:
                p = doc.add_paragraph(style='List Number')
                format_inline_text(p, item)
            numbered_list = []
            in_numbered_list = False
        
        if elem_type == 'checkbox':
            checked, text = content
            p = doc.add_paragraph()
            checkbox_char = '☑' if checked else '☐'
            run = p.add_run(f'{checkbox_char} ')
            run.font.size = Pt(12)
            format_inline_text(p, text)
            continue
        
        # Handle other elements
        if elem_type == 'empty':
            continue
        elif elem_type == 'hr':
            # Add horizontal line
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            # Add a border below
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), '003366')
            pBdr.append(bottom)
            p._p.get_or_add_pPr().append(pBdr)
        elif elem_type == 'h1':
            # Remove emoji from heading
            clean_text = re.sub(r'[^\x00-\x7F]+', '', content).strip()
            if clean_text:
                p = doc.add_heading(clean_text, level=1)
        elif elem_type == 'h2':
            clean_text = re.sub(r'[^\x00-\x7F]+', '', content).strip()
            if clean_text:
                p = doc.add_heading(clean_text, level=2)
        elif elem_type == 'h3':
            clean_text = re.sub(r'[^\x00-\x7F]+', '', content).strip()
            if clean_text:
                p = doc.add_heading(clean_text, level=3)
        elif elem_type == 'h4':
            clean_text = re.sub(r'[^\x00-\x7F]+', '', content).strip()
            if clean_text:
                p = doc.add_heading(clean_text, level=4)
        elif elem_type == 'table':
            add_table(doc, content)
        elif elem_type == 'code':
            for code_line in content.split('\n'):
                p = doc.add_paragraph(style='GrantCode')
                p.add_run(code_line)
        elif elem_type == 'quote':
            p = doc.add_paragraph(style='GrantQuote')
            format_inline_text(p, content)
        elif elem_type == 'bold_line':
            p = doc.add_paragraph()
            # Parse the bold line
            text = content
            if text.startswith('**'):
                match = re.match(r'\*\*([^*]+)\*\*(.*)', text)
                if match:
                    run = p.add_run(match.group(1))
                    run.bold = True
                    if match.group(2):
                        format_inline_text(p, match.group(2).lstrip())
                else:
                    format_inline_text(p, text)
            else:
                format_inline_text(p, text)
        elif elem_type == 'paragraph':
            p = doc.add_paragraph(style='GrantBody')
            format_inline_text(p, content)
    
    # Handle any remaining lists
    if bullet_list:
        for item in bullet_list:
            p = doc.add_paragraph(style='List Bullet')
            format_inline_text(p, item)
    
    if numbered_list:
        for item in numbered_list:
            p = doc.add_paragraph(style='List Number')
            format_inline_text(p, item)
    
    # Save document
    doc.save(docx_path)
    print(f"Created: {docx_path}")


def main():
    """Main function to convert all markdown files."""
    base_path = Path(r'C:\nanobot\instance3\workspace\grant_templates')
    
    # Define files to convert
    files_to_convert = [
        ('templates/united_way_full.md', 'templates/united_way_full.docx'),
        ('templates/united_way_canada_template.md', 'templates/united_way_canada_template.docx'),
        ('templates/federal_canada.md', 'templates/federal_canada.docx'),
        ('templates/foundation_private.md', 'templates/foundation_private.docx'),
        ('templates/corporate_sponsorship.md', 'templates/corporate_sponsorship.docx'),
        ('sections/organization_sections.md', 'sections/organization_sections.docx'),
        ('budgets/budget_templates.md', 'budgets/budget_templates.docx'),
        ('samples/youth_mentoring_sample.md', 'samples/youth_mentoring_sample.docx'),
        ('samples/sample_proposal.md', 'samples/sample_proposal.docx'),
        ('checklists/submission_checklists.md', 'checklists/submission_checklists.docx'),
        ('research/united_way_east_ontario.md', 'research/united_way_east_ontario.docx'),
        ('research/quick_reference.md', 'research/quick_reference.docx'),
        ('README.md', 'README.docx'),
    ]
    
    print("=" * 60)
    print("Grant Template Library: Markdown to Word Converter")
    print("=" * 60)
    print()
    
    success_count = 0
    error_count = 0
    
    for md_file, docx_file in files_to_convert:
        md_path = base_path / md_file
        docx_path = base_path / docx_file
        
        if md_path.exists():
            try:
                convert_md_to_docx(md_path, docx_path)
                success_count += 1
            except Exception as e:
                print(f"ERROR converting {md_file}: {e}")
                error_count += 1
        else:
            print(f"NOT FOUND: {md_file}")
            error_count += 1
    
    print()
    print("=" * 60)
    print(f"Conversion complete: {success_count} files converted, {error_count} errors")
    print("=" * 60)


if __name__ == '__main__':
    main()
